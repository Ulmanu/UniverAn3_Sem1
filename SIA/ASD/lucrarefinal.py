import tkinter as tk #interfata GUI
from tkinter import ttk
import requests #apelam siteul
import re #expresii regulate
from bs4 import BeautifulSoup #convertim html
from xlwt import Workbook #inscriem in excel
from docx import Document #inscriem in word
import os #cream fisiere si salvam in ele
from rake_nltk import Rake #extragem cuvintele cheie
import nltk #analizam textul
from os import path #stergem fisiere
from nltk.corpus import stopwords #gasim stopwords
import gethtml #modul creat de mine se foloseste la cautarea cheilor dupa url
#cream interfata grafica
form=tk.Tk()
form.title("Lucrare de An ASD")
form.geometry("1200x1000")
tab_parent=ttk.Notebook(form)
tab1=ttk.Frame(tab_parent)
tab2=ttk.Frame(tab_parent)
tab_parent.add(tab1, text="Extract Articles")
tab_parent.add(tab2, text="Extract Keywords")

#===First tab===
#===GUI====
label1=tk.Label(tab1, text="Keyword:")
entry1=tk.Entry(tab1)
text1=tk.Text(tab1, height=15, width=25)
label1.grid(row=0,column=0,padx=15,pady=15)
entry1.grid(row=0,column=1,padx=15,pady=15)
#functie de inscriere keywords in xls
def sheet(text):
    #adresa site-ului, adaugam xls 
    doclink="http://www.scholarpedia.org"+text
    print(doclink)
    wb = Workbook()
    sheet1 = wb.add_sheet('Sheet 1') 
    headers1 = {'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/55.0.2883.87 Safari/537.36'}
    document = Document()
    link1=requests.get(doclink,headers=headers1).text
    soup1=BeautifulSoup(link1,'lxml')
    #extragem info din tagul p si h1
    selectall2=soup1.find_all("p")
    selecttitle=str(*soup1.find_all("h1"))
    print(selecttitle)
    #print(*selectall,sep='\n')
    stri2 = ""
    i=0
    for lin in selectall2:
        stri2+=str(lin)
        i=i+1
    #filtram textul de taguri html
    clean = re.compile('<.*?>|&([a-z0-9]+|#[0-9]{1,6}|#x[0-9a-f]{1,6});')
    stri2=re.sub(clean, '', stri2)
    sheet1.write(0, 0, 'Cuvinte Cheie')
    sheet1.write(0, 1, 'Rake Rank')
    #utilizam algortimtul rake
    r=Rake()
    r.extract_keywords_from_text(stri2)
    print("\n".join(r.get_ranked_phrases()))
    print(*r.get_ranked_phrases_with_scores(),sep='\n')
    #rank,key=r.get_ranked_phrases_with_scores()
    rank=r.get_ranked_phrases_with_scores()
    print(rank)
    word=[0 for x in range(len(rank))]
    ranked=[0 for x in range(len(rank))]
    j=0
    for khh in rank:
        ranked[j],word[j]=khh
        j=j+1
    #inscriem in xls 
    for g in range(len(ranked)):
        sheet1.write(g+1, 0, word[g])
        sheet1.write(g+1, 1, ranked[g])
    if path.exists("key.xls"):
        os.remove("key.xls")
        wb.save('key.xls' )
    else:
        wb.save('key.xls' )
#functie de inscriere in document a articolelor
def document(text):
    #adresa site-ului, adaugam doc
    doclink="http://www.scholarpedia.org"+text
    print(doclink)
    headers1 = {'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/55.0.2883.87 Safari/537.36'}
    document = Document()
    link1=requests.get(doclink,headers=headers1).text
    soup1=BeautifulSoup(link1,'lxml')
    #extragem info din tagul p si h1
    selectall2=soup1.find_all("p")
    selecttitle=str(*soup1.find_all("h1"))
    print(selecttitle)
    #print(*selectall,sep='\n')
    stri2 = [''  for stri2 in range(len(selectall2))]
    i=0
    for lin in selectall2:
        stri2[i]=str(lin)
        i=i+1
    #filtram textul de taguri html
    clean = re.compile('<.*?>|&([a-z0-9]+|#[0-9]{1,6}|#x[0-9a-f]{1,6});')
    selecttitle=re.sub(clean, '', selecttitle)
    document.add_heading(selecttitle, level=1)
    for i in stri2:
        clean = re.compile('<.*?>')
        i=re.sub(clean, '', i)
        document.add_paragraph(i, style='Intense Quote')
    #inscriem in doc file
    if path.exists("document.docx"):
        os.remove("document.docx")
        document.save('document.docx' )
    else:
        document.save('document.docx' )
#funtie de inscriere a rezumatului in document
def documentsummary(text):
    doclink="http://www.scholarpedia.org"+text
    print(doclink)
    headers1 = {'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/55.0.2883.87 Safari/537.36'}
    document = Document()
    link1=requests.get(doclink,headers=headers1).text
    soup1=BeautifulSoup(link1,'lxml')
    selectall2=soup1.find_all("p")
    selecttitle=str(*soup1.find_all("h1"))
    print(selecttitle)
    #print(*selectall,sep='\n')
    stri2 = [''  for stri2 in range(len(selectall2))]
    i=0
    for lin in selectall2:
        stri2[i]=str(lin)
        i=i+1
    clean = re.compile('<.*?>|&([a-z0-9]+|#[0-9]{1,6}|#x[0-9a-f]{1,6});')
    selecttitle=re.sub(clean, '', selecttitle)
    document.add_heading(selecttitle, level=1)
    for i in range(len(selecttitle)):
        clean = re.compile('<.*?>|&([a-z0-9]+|#[0-9]{1,6}|#x[0-9a-f]{1,6});')
        stri2[i]=re.sub(clean, '', stri2[i])

    i=0
     #utilizarea algoritmului de rezumare a textului
    while i<len(stri2):
        text=str(stri2[i])
        print(text)
        stopWords = set(stopwords.words("english")) 
        words = nltk.word_tokenize(text) 

        freqTable = dict() 
        for word in words: 
                word = word.lower() 
                if word in stopWords: 
                    continue
                if word in freqTable: 
                    freqTable[word] += 1
                else: 
                    freqTable[word] = 1
        
     
        sentences = nltk.sent_tokenize(text) 
        sentenceValue = dict() 

        for sentence in sentences: 
            for word, freq in freqTable.items(): 
                if word in sentence.lower(): 
                    if sentence in sentenceValue:
                        sentenceValue[sentence] += freq 
                    else:
                        sentenceValue[sentence] = freq 

        sumValues = 0
        for sentence in sentenceValue: 
            sumValues += sentenceValue[sentence] 


        if len(sentenceValue)>0:
            average = int(sumValues / len(sentenceValue))
            summary = ''
            for sentence in sentences:
                if (sentence in sentenceValue) and (sentenceValue[sentence] > (1.2 * average)): 
                    summary = " " + sentence
                    print(summary)
                    clean = re.compile('<.*?>|&([a-z0-9]+|#[0-9]{1,6}|#x[0-9a-f]{1,6});')
                    summary=re.sub(clean, '', summary)
                    document.add_paragraph(summary, style='Intense Quote')
                
        else:
            document.add_paragraph(text, style='Intense Quote')
            
        i=i+1

    if path.exists("documentrez.docx"):
        os.remove("documentrez.docx")
        document.save('documentrez.docx' )
    else:
        document.save('documentrez.docx' )
    
#functie de afisare a rezultatelor cautarii dupa cuvant cheie
def titlecontent():
    query=entry1.get()
    who_link="http://www.scholarpedia.org/w/index.php?title=Special%3ASearch&profile=default&search="+query+"&fulltext=Search"
    headers = {'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/55.0.2883.87 Safari/537.36'}
    link=requests.get(who_link,headers=headers).text
    soup=BeautifulSoup(link,'lxml')
    selectall=soup.find_all("div",{"class": "searchresult"})
    #print(*selectall,sep='\n')
    stri = [''  for stri in range(len(selectall))]
    i=0
    for lin in selectall:
        stri[i]=str(lin)
        i=i+1
    for i in stri:
        clean = re.compile('<.*?>')
        i=re.sub(clean, '', i)
        
    selectall1=soup.find_all("div",{"class": "mw-search-result-heading"})
    #print(*selectall,sep='\n')

    stri1 = [''  for stri1 in range(len(selectall))]
    i=0
    for lin in selectall1:
        stri1[i]=str(lin)
        i=i+1

    for i in stri1:
        clean = re.compile('<.*?>')
        i=re.sub(clean, '', i)        

    kl = [''  for kl in range(len(selectall))]
    i=0
    for lin in selectall1:
        kl[i]=str(lin.find('a').get('href'))
        i=i+1
    
    btntxt=[0 for x in range(len(selectall))]
    btntxt1=[0 for x in range(len(selectall))]
    btntxt2=[0 for x in range(len(selectall))]
    txt=[0 for x in range(len(selectall))]  
    for x in range(20):
        txt[x] = tk.Text(tab1,height=2, width=35)
        clean = re.compile('<.*?>')
        stri1[x]=re.sub(clean, '', stri1[x])
        btntxt[x]=tk.Button(tab1, text="Document_"+str(x),command =lambda x1=x: document(kl[x1]))
        btntxt1[x]=tk.Button(tab1, text="Summarize_"+str(x),command =lambda x1=x: documentsummary(kl[x1]))
        btntxt2[x]=tk.Button(tab1, text="Keywords_"+str(x),command =lambda x1=x: sheet(kl[x1]))
        print(kl[x])
        txt[x].insert(tk.INSERT, stri1[x])
        txt[x].grid(column=0, row=x+1)
        btntxt[x].grid(column=2, row=x+1)
        btntxt1[x].grid(column=3, row=x+1)
        btntxt2[x].grid(column=4, row=x+1)

    txt1=[0 for x in range(len(selectall))]  
    for x in range(20):
        txt1[x] = tk.Text(tab1,height=2, width=80)
        clean = re.compile('<.*?>')
        stri[x]=re.sub(clean, '', stri[x])
        txt1[x].insert(tk.INSERT, stri[x])
        txt1[x].grid(column=1, row=x+1)

#functie de inscriere a rezultatelor cautarii in xls
def excel():
    query=entry1.get()
    wb = Workbook()
    sheet1 = wb.add_sheet('Sheet 1') 
    who_link="http://www.scholarpedia.org/w/index.php?title=Special%3ASearch&profile=default&search="+query+"&fulltext=Search"
    headers = {'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/55.0.2883.87 Safari/537.36'}
    link=requests.get(who_link,headers=headers).text
    soup=BeautifulSoup(link,'lxml')
    selectall=soup.find_all("div",{"class": "searchresult"})
    #print(*selectall,sep='\n')
    stri = [''  for stri in range(len(selectall))]
    i=0
    for lin in selectall:
        stri[i]=str(lin)
        i=i+1

    for i in stri:
        clean = re.compile('<.*?>')
        i=re.sub(clean, '', i)
        

    selectall1=soup.find_all("div",{"class": "mw-search-result-heading"})
    #print(*selectall,sep='\n')

    stri1 = [''  for stri1 in range(len(selectall))]
    i=0
    for lin in selectall1:
        stri1[i]=str(lin)
        i=i+1


    for i in stri1:
        clean = re.compile('<.*?>')
        i=re.sub(clean, '', i)
        
     
    for x in range(20):
        
        clean = re.compile('<.*?>')
        stri1[x]=re.sub(clean, '', stri1[x])
        
   
    for x in range(20):
        
        clean = re.compile('<.*?>')
        stri[x]=re.sub(clean, '', stri[x])

    sheet1.write(0, 0, 'DENUMIRI ARTICOLE')
    sheet1.write(0, 1, 'DESCRIERE ARTICOLE')
    for x in range(20):
        sheet1.write(x+1, 0, stri1[x])
        sheet1.write(x+1, 1, stri[x])

    wb.save(query+'.xls') 

btn1=tk.Button(tab1, text="Search",command = titlecontent)
btn1.grid(row=0,column=2,padx=15,pady=15)

btn2=tk.Button(tab1, text="Excel Save",command = excel)
btn2.grid(row=0,column=3,padx=15,pady=15)


#===Second tab===
#extragerea keywords din site dupa url
def keywordextr():
    text21.delete("1.0",tk.END)
    query=entry21.get()
    print(gethtml.gethtml(query))
    article=gethtml.getarticle(query)
    word=gethtml.getkeywords(article)
    index=gethtml.getkeywordsindex(article)
    for w in range(30):
        text21.insert(tk.INSERT, str(word[w])+" "+str(index[w])+"\n")
        
    print(gethtml.getkeywords(article))
#inscrierea keywords in excel
def excelfinal():
    wb = Workbook()
    sheet1 = wb.add_sheet('Sheet 1')
    sheet1.write(0, 0, 'Cuvinte Cheie')
    sheet1.write(0, 1, 'Nr. aparitii')
    query=entry21.get()
    print(gethtml.gethtml(query))
    article=gethtml.getarticle(query)
    word=gethtml.getkeywords(article)
    index=gethtml.getkeywordsindex(article)
    for w in range(30):
        sheet1.write(w+1, 0, word[w])
        sheet1.write(w+1, 1, index[w])
        
    if path.exists("keywords.xls"):
        os.remove("keywords.xls")
        wb.save('keywords.xls' )
    else:
        wb.save('keywords.xls' )
        

        
label21=tk.Label(tab2, text="Site:")
entry21=tk.Entry(tab2,width=50)
text21=tk.Text(tab2, height=30, width=20)
label21.grid(row=0,column=0,padx=15,pady=15)
entry21.grid(row=0,column=1,padx=15,pady=15)
btn21=tk.Button(tab2, text="Extract keywords",command = keywordextr)
btn21.grid(row=0,column=2,padx=15,pady=15)
text21.grid(row=1,column=0,padx=15,pady=15)
btn22=tk.Button(tab2, text="Excel Save",command = excelfinal)
btn22.grid(row=0,column=3,padx=15,pady=15)
tab_parent.pack(expand=1, fill='both')

form.mainloop()
